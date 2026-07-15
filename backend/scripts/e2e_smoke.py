"""End-to-end smoke test for a running local API server.

Usage:
    E2E_API_URL=http://127.0.0.1:8002 python scripts/e2e_smoke.py
"""
from __future__ import annotations

import asyncio
import io
import json
import os
import time
from typing import Any

import httpx


API_URL = os.getenv("E2E_API_URL", "http://127.0.0.1:8000").rstrip("/")


def _make_docx_bytes() -> bytes:
    """Create a small real DOCX fixture so the smoke test exercises parsing."""
    from docx import Document

    doc = Document()
    doc.add_heading("E2E Biology Lecture", level=1)
    doc.add_paragraph(
        "Photosynthesis converts light energy into chemical energy through "
        "chloroplast reactions, while cellular respiration releases stored "
        "energy through glycolysis, the Krebs cycle, and oxidative phosphorylation."
    )
    doc.add_paragraph(
        "Mitosis and meiosis are central exam concepts because they compare "
        "cell division outcomes, chromosome behavior, and genetic variation."
    )
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _sse_events(text: str) -> list[dict[str, Any]]:
    events: list[dict[str, Any]] = []
    for line in text.splitlines():
        if not line.startswith("data: "):
            continue
        events.append(json.loads(line.removeprefix("data: ")))
    return events


def _smoke_answer_for(question: dict[str, Any]) -> str:
    """Return a non-blank answer that satisfies the submit validation contract."""
    question_type = question.get("question_type")
    choices = question.get("choices") or []
    if question_type == "multiple_choice" and choices:
        return str(choices[0].get("label") or choices[0].get("text") or "A")
    if question_type == "true_false":
        return "True"
    return (
        "Photosynthesis, respiration, and cell division are connected exam "
        "concepts because they explain how cells transform energy and pass "
        "genetic information."
    )


async def _wait_for_material(
    client: httpx.AsyncClient,
    course_id: str,
    token: str,
    filename: str,
    wanted: set[str],
) -> dict[str, Any]:
    headers = {"Authorization": f"Bearer {token}"}
    for _ in range(30):
        res = await client.get(f"/courses/{course_id}/materials", headers=headers)
        res.raise_for_status()
        for material in res.json():
            if material["original_filename"] == filename and material["processing_status"] in wanted:
                return material
        await asyncio.sleep(0.5)
    raise RuntimeError(f"Timed out waiting for {filename} to reach {sorted(wanted)}")


async def main() -> None:
    async with httpx.AsyncClient(base_url=API_URL, timeout=60.0) as client:
        email = f"e2e-{int(time.time())}@example.com"
        password = "password123"
        token: str | None = None
        registered = False
        cleanup_completed = False

        try:
            health = await client.get("/health")
            health.raise_for_status()

            register = await client.post(
                "/auth/register",
                json={
                    "email": email,
                    "password": password,
                    "full_name": "E2E Student",
                },
            )
            register.raise_for_status()
            registered = True

            login = await client.post(
                "/auth/login",
                data={"username": email, "password": password},
            )
            login.raise_for_status()
            token = login.json()["access_token"]
            headers = {"Authorization": f"Bearer {token}"}

            course = await client.post(
                "/courses",
                json={
                    "name": "E2E Biology",
                    "description": "Smoke test course",
                    "professor_name": "Dr. Smoke",
                    "subject": "Biology",
                },
                headers=headers,
            )
            course.raise_for_status()
            course_id = course.json()["id"]

            upload = await client.post(
                f"/courses/{course_id}/materials",
                files={
                    "files": (
                        "lecture.docx",
                        _make_docx_bytes(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
                headers=headers,
            )
            upload.raise_for_status()
            await _wait_for_material(client, course_id, token, "lecture.docx", {"completed"})

            bad_upload = await client.post(
                f"/courses/{course_id}/materials",
                files={"files": ("bad.png", b"not an image", "image/png")},
                headers=headers,
            )
            bad_upload.raise_for_status()
            failed = await _wait_for_material(client, course_id, token, "bad.png", {"failed"})
            retry = await client.post(
                f"/courses/{course_id}/materials/{failed['id']}/retry",
                headers=headers,
            )
            retry.raise_for_status()
            retried_material = await _wait_for_material(
                client,
                course_id,
                token,
                "bad.png",
                {"completed", "failed"},
            )

            analysis = await client.post(f"/courses/{course_id}/analysis", headers=headers)
            analysis.raise_for_status()
            analysis_events = _sse_events(analysis.text)
            if not any(event.get("type") == "complete" for event in analysis_events):
                raise RuntimeError("Analysis stream did not complete")

            exam_stream = await client.post(
                f"/courses/{course_id}/exams",
                json={"title": "E2E Practice", "question_count": 3, "mode": "standard"},
                headers=headers,
            )
            exam_stream.raise_for_status()
            exam_events = _sse_events(exam_stream.text)
            complete = next((event for event in exam_events if event.get("type") == "complete"), None)
            if not complete or not complete.get("exam_id"):
                raise RuntimeError("Exam stream did not return an exam_id")
            exam_id = complete["exam_id"]

            exam = await client.get(f"/exams/{exam_id}", headers=headers)
            exam.raise_for_status()
            questions = exam.json()["questions"]
            answers = [
                {"question_id": question["id"], "student_answer": _smoke_answer_for(question)}
                for question in questions
            ]
            submit = await client.post(
                f"/exams/{exam_id}/submit",
                json={"answers": answers},
                headers=headers,
            )
            submit.raise_for_status()

            result = await client.get(f"/exams/{exam_id}/result", headers=headers)
            result.raise_for_status()

            summary = {
                "email": email,
                "course_id": course_id,
                "exam_id": exam_id,
                "score": result.json()["score"],
                "questions": len(result.json()["results"]),
                "material_retry_status": retried_material["processing_status"],
            }

            # Smoke data should never accumulate in a developer's local database.
            cleanup = await client.delete("/auth/me", headers=headers)
            cleanup.raise_for_status()
            cleanup_completed = True
            summary["cleanup"] = "account_deleted"

            print(json.dumps(summary, indent=2))
        finally:
            if registered and not cleanup_completed:
                try:
                    if token is None:
                        login = await client.post(
                            "/auth/login",
                            data={"username": email, "password": password},
                        )
                        if login.is_success:
                            token = login.json()["access_token"]
                    if token is not None:
                        cleanup = await client.delete(
                            "/auth/me",
                            headers={"Authorization": f"Bearer {token}"},
                        )
                        if cleanup.status_code != 204:
                            print(
                                f"Failed to clean up smoke account {email}: HTTP {cleanup.status_code}",
                                file=sys.stderr,
                            )
                except httpx.HTTPError as exc:
                    print(f"Failed to clean up smoke account {email}: {exc}", file=sys.stderr)


if __name__ == "__main__":
    asyncio.run(main())
